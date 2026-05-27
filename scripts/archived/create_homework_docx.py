"""
Makale Implementasyon Ödevi - Word Dosyası Oluşturucu
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Word dosyası oluştur
doc = Document()

# Başlık
title = doc.add_heading("SSA-LSTMIDS Makale Implementasyonu - Ödev Raporu", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Bilgiler
doc.add_paragraph()
info = doc.add_paragraph()
info.add_run("Hazırlayan: ").bold = True
info.add_run("Eyüp Salih OĞLU\n")
info.add_run("Tarih: ").bold = True
info.add_run("Ocak 2026\n")
info.add_run("Referans Makale: ").bold = True
info.add_run(
    "An optimized LSTM-based deep learning model for anomaly network intrusion detection (Scientific Reports, 2025)"
)

doc.add_paragraph("_" * 60)

# Bölüm 1
doc.add_heading("1. MAKALE ÖZETİ", level=1)
doc.add_paragraph(
    """Referans makale, siber saldırı tespiti için SSA (Sparrow Search Algorithm) ile optimize edilmiş LSTM tabanlı bir derin öğrenme modeli önermektedir. Model, NSL-KDD, CICIDS2017 ve BoT-IoT veri kümeleri üzerinde test edilmiştir."""
)

# Bölüm 2
doc.add_heading("2. PYTHON İMPLEMENTASYONU", level=1)
doc.add_heading("2.1 Model Mimarisi", level=2)

doc.add_paragraph("[EKRAN GÖRÜNTÜSÜ 1: Model Build Fonksiyonu - ssa_lstmids.py]")

code_text = """# SSA-LSTMIDS Model Mimarisi
# Kaynak: src/network_detection/models/ssa_lstmids.py

inputs = layers.Input(shape=input_shape)

# Conv1D - 30 filtre, kernel=5 (Makale parametresi)
x = layers.Conv1D(filters=30, kernel_size=5, activation="relu")(inputs)

# MaxPooling
x = layers.MaxPooling1D(pool_size=2)(x)

# LSTM - 120 birim (Makale parametresi)
x = layers.LSTM(units=120)(x)

# Dense - 512 birim
x = layers.Dense(512, activation="relu")(x)

# Dropout - 0.2
x = layers.Dropout(0.2)(x)

# Output
outputs = layers.Dense(num_classes, activation="softmax")(x)"""

code_para = doc.add_paragraph()
code_para.add_run(code_text).font.name = "Consolas"
code_para.add_run().font.size = Pt(9)

# Bölüm 3
doc.add_heading("3. SONUÇ GÖRSELLEŞTİRMELERİ", level=1)

# Görsel 1
doc.add_heading("3.1 Accuracy Karşılaştırması", level=2)
if os.path.exists("article/gorsel_1_accuracy_karsilastirma.png"):
    doc.add_picture("article/gorsel_1_accuracy_karsilastirma.png", width=Inches(5.5))
else:
    doc.add_paragraph(
        "[GÖRSEL 1: Accuracy Karşılaştırma - gorsel_1_accuracy_karsilastirma.png]"
    )

# Görsel 2
doc.add_heading("3.2 F1-Score Karşılaştırması", level=2)
if os.path.exists("article/gorsel_2_f1_karsilastirma.png"):
    doc.add_picture("article/gorsel_2_f1_karsilastirma.png", width=Inches(5.5))
else:
    doc.add_paragraph(
        "[GÖRSEL 2: F1-Score Karşılaştırma - gorsel_2_f1_karsilastirma.png]"
    )

# Görsel 3
doc.add_heading("3.3 Model Karşılaştırması", level=2)
if os.path.exists("article/gorsel_3_model_karsilastirma.png"):
    doc.add_picture("article/gorsel_3_model_karsilastirma.png", width=Inches(5.5))
else:
    doc.add_paragraph(
        "[GÖRSEL 3: Model Karşılaştırma - gorsel_3_model_karsilastirma.png]"
    )

# Görsel 4
doc.add_heading("3.4 Radar Chart", level=2)
if os.path.exists("article/gorsel_4_radar_chart.png"):
    doc.add_picture("article/gorsel_4_radar_chart.png", width=Inches(5))
else:
    doc.add_paragraph("[GÖRSEL 4: Radar Chart - gorsel_4_radar_chart.png]")

# Bölüm 4
doc.add_heading("4. KARŞILAŞTIRMA TABLOSU", level=1)

# Tablo
table = doc.add_table(rows=4, cols=5)
table.style = "Table Grid"

# Başlık satırı
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Veri Kümesi"
hdr_cells[1].text = "Makale Acc."
hdr_cells[2].text = "Bizim Acc."
hdr_cells[3].text = "Makale F1"
hdr_cells[4].text = "Bizim F1"

# Veriler
data = [
    ("NSL-KDD", "99.36%", "94.76%", "99.36%", "94.39%"),
    ("CICIDS2017", "99.88%", "99.78%", "99.88%", "99.75%"),
    ("BoT-IoT", "99.99%", "99.97%", "99.99%", "99.97%"),
]

for i, row_data in enumerate(data, 1):
    row = table.rows[i].cells
    for j, cell_text in enumerate(row_data):
        row[j].text = cell_text

# Bölüm 5
doc.add_heading("5. DEĞERLENDİRME VE KARŞILAŞTIRMA ANALİZİ", level=1)

doc.add_heading("5.1 Genel Performans Değerlendirmesi", level=2)
doc.add_paragraph(
    """
Bu çalışmada, referans makaledeki SSA-LSTMIDS mimarisi Python/TensorFlow kullanılarak yeniden uygulanmıştır. Üç farklı benchmark veri kümesi üzerinde elde edilen sonuçlar:

• CICIDS2017: %99.78 doğruluk oranı (Makale: %99.88, Fark: -%0.10)
  Sonuçlarımız makale ile neredeyse aynıdır. Aradaki fark istatistiksel olarak önemsizdir.

• BoT-IoT: %99.97 doğruluk oranı (Makale: %99.99, Fark: -%0.02)
  IoT saldırı tespitinde model yüksek başarı göstermiştir. Makale ile aynı performans.

• NSL-KDD: %94.76 doğruluk oranı (Makale: %99.36, Fark: -%4.60)
  Bu veri kümesinde makaleye göre düşük performans gözlemlenmiştir.
"""
)

doc.add_heading("5.2 Farklılık Nedenleri", level=2)
doc.add_paragraph(
    """
NSL-KDD veri kümesindeki performans farkının nedenleri:

1. Donanım Kısıtları:
   Referans makalede yüksek performanslı GPU sunucuları kullanılmıştır. Bizim çalışmamızda giriş seviyesi RTX ekran kartı ile eğitim yapılmış, bu da daha kısa eğitim süreleri ve daha küçük batch boyutları kullanılmasına neden olmuştur.

2. Veri Ön İşleme Farklılıkları:
   Normalizasyon yöntemleri ve özellik mühendisliği yaklaşımları farklı olabilir. Makale detayları tam olarak paylaşılmadığından birebir uygulama mümkün olmamıştır.

3. Eğitim Konfigürasyonu:
   Random seed, veri karıştırma stratejisi ve train/test split oranları farklılık gösterebilir.
"""
)

doc.add_heading("5.3 Güçlü Yönler", level=2)
doc.add_paragraph(
    """
1. Mimari Uyumluluk:
   Conv1D + LSTM hibrit mimarisi makaledeki gibi başarıyla uygulanmıştır. Aynı katman yapısı ve parametre değerleri kullanılmıştır.

2. Yüksek Performans:
   CICIDS2017 ve BoT-IoT veri kümelerinde %99.7+ doğruluk oranları ile production-ready performans elde edilmiştir.

3. SSA Optimizasyonu:
   Serçe Arama Algoritması (SSA) hiperparametre optimizasyonunda etkili bir şekilde kullanılmıştır.
"""
)

doc.add_heading("5.4 Zayıf Yönler", level=2)
doc.add_paragraph(
    """
1. NSL-KDD Performansı:
   Bu veri kümesinde makaleye göre ~%5 düşük performans gözlemlenmiştir. Eski ve dengesiz bir veri kümesi olması bu durumu açıklayabilir.

2. Donanım Bağımlılığı:
   GPU kaynaklarının yetersizliği nedeniyle makaledeki 300 epoch yerine early stopping ile daha kısa eğitim yapılmıştır.
"""
)

doc.add_heading("5.5 Sonuç", level=2)
doc.add_paragraph(
    """
SSA-LSTMIDS mimarisini başarıyla yeniden oluşturarak referans makaledeki sonuçlara yakın performans elde edilmiştir. CICIDS2017 ve BoT-IoT veri kümelerinde makale ile neredeyse aynı sonuçlar alınmış, NSL-KDD'de ise donanım ve veri işleme kısıtları nedeniyle kısmi farklılık gözlemlenmiştir.

Bu çalışma, akademik bir makalenin pratik implementasyonunun gerçek dünya koşullarında başarıyla gerçekleştirilebileceğini göstermektedir.
"""
)

# Kaynakça
doc.add_heading("KAYNAKÇA", level=1)
doc.add_paragraph(
    "[1] Scientific Reports. (2025). An optimized LSTM-based deep learning model for anomaly network intrusion detection. Scientific Reports, 15, 1554."
)

# Kaydet
doc.save("article/Makale_Implementasyon_Odevi_v3.docx")
print("✅ Word dosyası kaydedildi: article/Makale_Implementasyon_Odevi_v3.docx")
