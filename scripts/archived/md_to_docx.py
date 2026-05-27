"""
Markdown to DOCX converter for CyberGuard AI paper
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os


def md_to_docx(md_path: str, docx_path: str, new_title: str = None):
    """Convert markdown file to DOCX"""

    # Read markdown
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Process lines
    lines = content.split("\n")
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    for line in lines:
        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_content)
                p = doc.add_paragraph()
                p.style = "Normal"
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Tables
        if line.strip().startswith("|") and "|" in line[1:]:
            if not in_table:
                in_table = True
                table_rows = []

            # Skip separator lines
            if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                continue

            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            continue
        elif in_table:
            # End table
            if table_rows:
                cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=cols)
                table.style = "Table Grid"
                for i, row_data in enumerate(table_rows):
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = cell_text
                doc.add_paragraph()
            in_table = False
            table_rows = []

        # Skip empty lines but add paragraph break
        if not line.strip():
            continue

        # Headers
        if line.startswith("# "):
            text = line[2:].strip()
            if new_title and "CyberGuard AI:" in text:
                text = new_title
            p = doc.add_heading(text, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)

        # Bold text like **text**
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True

        # Regular paragraph
        elif not line.startswith("---") and not line.startswith("```"):
            # Clean markdown formatting
            text = line
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # Bold
            text = re.sub(r"\*(.+?)\*", r"\1", text)  # Italic
            text = re.sub(r"`(.+?)`", r"\1", text)  # Code
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # Links

            if text.strip():
                doc.add_paragraph(text)

    # Handle remaining table
    if in_table and table_rows:
        cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                if j < cols:
                    table.rows[i].cells[j].text = cell_text

    # Save
    doc.save(docx_path)
    print(f"✅ DOCX created: {docx_path}")
    return docx_path


if __name__ == "__main__":
    import sys

    md_file = "article/CyberGuard_AI_Makale.md"
    docx_file = "article/CyberGuard_AI_Makale.docx"

    # Get new title from command line if provided
    new_title = sys.argv[1] if len(sys.argv) > 1 else None

    md_to_docx(md_file, docx_file, new_title)
